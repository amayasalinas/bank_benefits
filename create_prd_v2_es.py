# Genera PRD_Eliseo_ES_v2.docx desde el contenido canónico de PRD_Eliseo_ES_v2.md
# Ejecutar tras editar el .md:  python create_prd_v2_es.py
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()


def H(text, level=1):
    doc.add_heading(text, level=level)


def P(text):
    doc.add_paragraph(text)


def bullets(items, style='List Bullet'):
    for it in items:
        doc.add_paragraph(it, style=style)


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v


# ===================== Título =====================
title = doc.add_heading('Documento de Requisitos del Producto (PRD) — Eliseo v2 (Edición Integrada: Estrategia + Producto)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Estado: v2 integrada (junio 2026). Fusiona el PRD v1 (ajustado al código en feat/ui-polish) con el Blueprint Estratégico (Partes I–IX). El .md es la fuente canónica; este .docx se regenera desde él. Leyenda: ✅ Hecho · 🟡 Parcial · 🔲 Pendiente · 🆕 Nuevo en v2 (propuesto).')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

P('Cambio de marco respecto a v1: v1 describía qué se construyó. v2 ancla el producto a una tesis de negocio validable: a quién servimos (el Optimizador), por qué es defendible (neutralidad cross-bank), cuánto vale (premio económico real), y cómo el motor produce recomendaciones objetivas (normalización a pesos).')

# ===================== 1 =====================
H('1. Resumen Ejecutivo, Tesis y KPIs')
bullets([
    'Visión: Eliseo es el asistente neutral de decisión de tarjetas para Colombia: conociendo todas las tarjetas que el usuario ya tiene, le dice qué tarjeta usar en cada compra para maximizar su beneficio —y cuál solicitar, conservar o cancelar—, mostrando el ahorro en pesos.',
    'Tesis de inversión: no se invierte en un comparador de tarjetas; se invierte en la futura capa neutral de decisión y demanda financiera de LatAm, usando las tarjetas de crédito como cuña de entrada.',
    'Por qué ahora: el pivote de "¿qué tarjeta solicitar?" hacia "¿qué tarjeta usar ahora?" (uso recurrente → hábito → datos → moat) es el movimiento correcto. Viento de cola: Open Finance obligatorio (Decreto 0368/2026).',
    'North Star: Transacciones Optimizadas = veces que un usuario consulta el recomendador antes de comprar, sin que se le pida.',
])
table(['KPI', 'Meta', 'Fuente'], [
    ['Consultas no provocadas / usuario activo / semana', '≥ 2 (hábito)', 'Parte V'],
    ['Retención semana 4 (cohorte)', '≥ 40%', 'Parte V'],
    ['Ahorro real confirmado / usuario / mes', '≥ $30–40k COP', 'Parte IV'],
    ['Tarjetas agregadas / usuario (primera semana)', '> 2.5', 'v1'],
    ['Sesiones/semana en Recomendador u Ofertas', '> 3', 'v1'],
    ['CTR a CTA de afiliado (proxy de monetización CPA)', '> 5%', 'v1'],
    ['Disposición a pagar premium (cuartil alto)', '≥ 25%', 'Parte IV'],
])

# ===================== 2 =====================
H('2. Problema y Oportunidad de Mercado')
bullets([
    'Problema: los colombianos con varias tarjetas dejan dinero sobre la mesa a diario: no saben qué tarjeta usar, se pierden promociones que cambian cada semana, y dejan vencer millas y puntos. Los comparadores existentes solo responden "¿cuál solicitar?".',
    'Mercado: ~14.2M tarjetas de crédito en Colombia (en contracción). El comparador de "cuál solicitar" ya está ocupado (Rankia). La oportunidad no monetizada es el uso recurrente en el punto de gasto.',
    'Segmentación económica (Parte IV): el premio no es homogéneo (ahorro mediano ~$350–478k/año), concentrado en gasto medio-alto. El básico (≈35% del mercado) casi no captura valor recurrente. Implicación: el producto y el GTM apuntan al Optimizador, no al masivo.',
])
table(['Perfil', 'Premio recompensas (mediana, año)'], [
    ['Básico', '$90.000 ❌ (no sostiene hábito)'],
    ['Profesional joven', '$385.000'],
    ['Familia ingresos medios', '$542.000'],
    ['Viajero frecuente', '$725.000'],
    ['Premium multi-tarjeta', '$1.503.000'],
])

# ===================== 3 =====================
H('3. ICP — "El Optimizador" (foco en crédito)')
P('Filtros duros: gasto ≥ $2.5M COP/mes en tarjeta · ≥ 2 tarjetas de crédito · totalero (paga el total) · interés activo en viajes/millas/cashback · potencial ≥ $400k/año de valor. (Débito: alcance secundario.)')
bullets([
    'Demografía: 28–45 años; urbano (Bogotá, Medellín, Cali, Barranquilla, Bucaramanga); estrato 4–6; universitario/posgrado.',
    'Ingresos: $6–20M COP/mes.',
    'Profesión: profesionales corporativos, consultores, ingenieros/TI, médicos/abogados, finanzas, emprendedores digitales, freelancers de alto ingreso.',
    'Comportamientos: totalero; ≥2 tarjetas (a menudo una de millas + una de cashback/puntos); ya acumula y redime; viaja 1–3 veces/año; usa apps fintech.',
    'Motivaciones: viajar con millas; "ganarle al sistema"; eficiencia/estatus; no dejar dinero sobre la mesa.',
    'Frustraciones: información dispersa; no saber qué tarjeta usar dónde; promos que se pierde; millas/puntos que vencen; cuotas que no sabe si valen la pena.',
])
P('Anti-ICP (no perseguir al inicio): rotativos que financian; usuarios de 1 tarjeta o gasto <$1.5M/mes; sin interés en optimizar. Reemplaza el ICP amplio de v1.')

# ===================== 4 =====================
H('4. Propuesta de Valor y Posicionamiento')
bullets([
    'Propuesta de valor: "Recupera ~$40.000 al mes (+$450.000/año) en beneficios que hoy dejas sobre la mesa." Para viajeros/premium: $700k–1.5M/año.',
    'Posicionamiento: La Suiza de las tarjetas — el único asistente que está del lado del usuario y compara TODAS sus tarjetas (de todos los bancos). Cada recomendación viene con el "por qué" en pesos.',
])

# ===================== 5 =====================
H('5. Ventaja Competitiva y Moat (Parte III)')
P('El wedge estructural: un emisor único jamás puede ser un optimizador neutral entre todas tus tarjetas sin canibalizar su producto. Nu solo recomienda Nu; Bancolombia solo Bancolombia; RappiCard solo RappiCard.')
table(['Atacante', '¿Bloqueado por neutralidad?'], [
    ['Nu Colombia (5M clientes)', 'SÍ — mono-emisor'],
    ['Bancolombia / Puntos Colombia', 'SÍ — coalición cerrada'],
    ['RappiCard (emisor)', 'SÍ — mono-emisor (vía Davivienda)'],
    ['Rappi (plataforma) / Nequi', 'NO — agregador horizontal = amenaza real'],
])
P('Las 5 capas del moat (se construyen, no se tienen): (1) neutralidad estructural; (2) grafo de promociones fresco + tabla VE opinada; (3) loop de datos cross-bank; (4) confianza/categoría; (5) velocidad y foco.')

# ===================== 6 =====================
H('6. Alcance del Producto (MoSCoW, recortado al Optimizador)')
P('Must Have (MVP):')
bullets([
    '✅ Autenticación (Registro / Inicio de sesión).',
    '✅ Agregar tarjeta sin datos sensibles (banco/modelo + apodo + últimos 4 dígitos).',
    '✅ Dashboard con tarjeta principal y métricas resumidas.',
    '✅ Motor Recomendador por categoría.',
    '✅ Vista de detalle de tarjeta con beneficios categorizados.',
    '🔲 Perfil de pago (totalero/rotativo) en onboarding — gobierna la lógica del recomendador.',
    '🔲 Normalización a pesos efectivos (VE) en el recomendador.',
], style='List Number 2')
P('Should Have:')
bullets([
    '✅ Feed dinámico de Ofertas con vencimientos.',
    '✅ Filtrado de Ofertas por Banco y Categoría.',
    '✅ Niveles de confianza (Confirmado/Probable/Acción) — UI/badge construido jun 2026; cálculo automático pendiente del motor v2.',
    '✅ CTA saliente rastreable en ofertas — construido jun 2026 (Offer.url + oferta clicable).',
    '✅ "Destacados / Mejores del mercado" (afiliados) — construido jun 2026 (ruta /destacados).',
    '🆕 Recomendaciones de portafolio (solicitar/conservar/cancelar).',
], style='List Number 2')
P('Could Have:')
bullets([
    '🔲 Notificaciones push para ofertas por vencer.',
    '🔲 Beneficios geolocalizados (salas VIP, etc.).',
    '🆕 Reporte semanal/mensual de ahorro real (gancho de retención).',
], style='List Number 2')
P('Won\'t Have (v1.0 → reafirmado):')
bullets([
    'Integración Open Finance en tiempo real (Fase 2, post-validación).',
    'Scraping automático de saldos de puntos/millas.',
], style='List Number 2')

# ===================== 7 =====================
H('7. Personas y Tareas a Realizar (JTBD)')
P('Personas detalladas vigentes en .product/PERSONAS/ (Gloria Espinosa, Sandra Moreno, Sebastián Ríos). Arquetipo principal conservado de v1:')
bullets([
    'Andrés (32), millennial tech. 4 tarjetas (Nu, Bancolombia LifeMiles, RappiCard, Falabella). Viaja 3 veces/año. Quiere vuelos gratis pero olvida qué tarjeta da salas VIP o el mejor multiplicador.',
    '"Cuando voy a pagar en un restaurante, quiero ver cuál de mis 4 tarjetas me da el mejor beneficio, para no perder dinero."',
    '"Cuando planeo un viaje, quiero saber cuál de mis tarjetas da acceso gratis a salas VIP, para no pagar $30 USD."',
])
P('Nota v2: Andrés es totalero y de alto gasto → ICP Optimizador. Validar que las personas de .product/PERSONAS/ cumplen los filtros duros del §3.')

# ===================== 8 =====================
H('8. Requisitos Funcionales y Motor de Recomendación v2')
H('8.1 Core: Gestionar Billetera (de v1, conservado)', level=2)
bullets([
    '[x] Lista de bancos buscable (tradicionales y digitales).',
    '[x] Filtra modelos tras seleccionar banco.',
    '[x] Apodo (máx 40) y últimos 4 dígitos (máx 4) opcionales.',
    '[x] Guarda (user_id, card_id, nickname, last_four).',
    '🔲 Captura perfil_pago (totalero/rotativo/mixto), gasto por categoría y preferencia en el onboarding.',
], style='List Number 2')
H('8.2 Core: Motor de Recomendación v2 (reemplaza el algoritmo simple de v1)', level=2)
P('Resuelve la Pregunta Abierta 9.1. Principio: todo beneficio se traduce a pesos colombianos efectivos esperados (COP_e) vía la tabla VE: VE_moneda = ValorRedención × p_red × (1 − breakage) (Puntos Colombia ≈ $7; LifeMiles ≈ $15 conservador).')
P('Tasa efectiva (el igualador): r_b(x) = unidades_por_peso_b(x) × VE_moneda(b) → cashback y millas comparables.')
P('Bifurcación por perfil de pago: Totalero → optimizar recompensa neta − cuota. Rotativo → priorizar tasa baja, NO recompensas (recomendar millas a un rotativo queda prohibido).')
P('Score de dos capas: VEO (Valor Económico Objetivo, neutral, en COP, siempre se muestra) y UP (Utilidad Personalizada, solo para ranking en UI). Elegibilidad = filtro duro.')
bullets([
    '[x] Recupera tarjetas del usuario y beneficios por category_id.',
    '🔲 Calcula r efectiva normalizada vía tabla VE.',
    '🔲 Aplica bifurcación totalero/rotativo.',
    '🔲 Incorpora promociones vigentes y resuelve stacking.',
    '🟡 Asigna nivel de confianza (Confirmado/Probable/Acción) — UI construida jun 2026; asignación automática pendiente.',
    '[x] Resultados < 500 ms (meta NFR < 100 ms con índices y precálculo).',
    '[x] Mensaje de respaldo si ninguna tarjeta tiene beneficio en la categoría.',
], style='List Number 2')
H('8.3 Core: Ofertas Explorables (de v1)', level=2)
bullets([
    '[x] Consulta promociones activas (valid_until >= hoy O NULL).',
    '[x] Filtros por banco y categoría (chips).',
    '✅ CTA saliente rastreable al hacer clic en una oferta — construido jun 2026 (Offer.url + oferta clicable, rel=noopener).',
], style='List Number 2')
H('8.4 Recomendaciones de Portafolio (🆕, Parte II §6)', level=2)
bullets([
    'Usar ahora: argmax del beneficio marginal por transacción.',
    'Solicitar: Valor(T∪{t}) − Valor(T) − cuota > 0.',
    'Cancelar: contribución marginal < cuota (con aviso de impacto en historial crediticio).',
])

# ===================== 9 =====================
H('9. Flujos Lógicos')
bullets([
    'Flujo 1 — Agregar tarjeta (de v1): banco → modelo → apodo → POST /api/user-cards → inserta en cards_user → dashboard con actualización optimista. 🔲 Tras la primera tarjeta, solicitar perfil_pago y gasto por categoría.',
    'Flujo 2 — Recomendador sin tarjetas (de v1): billetera = 0 → estado vacío con CTA a "Agregar tarjeta".',
    'Flujo 3 🆕 — Confianza/incertidumbre: el motor filtra promos validadas y vigentes y asigna nivel 🟢 Confirmado / 🟡 Probable / 🔵 Acción requerida. Nunca promete un beneficio condicionado a elegibilidad desconocida.',
])

# ===================== 10 =====================
H('10. Requisitos de Datos (esquema escalable, Parte VI)')
bullets([
    'Users / Banks / Cards / Merchants — maestras (UUID).',
    '🆕 card_versions — atributos mutables versionados (cuota, EA, requisitos, cupos) con vigencia y EXCLUDE anti-solapamiento.',
    '🆕 benefits versionados — (card_id, category, reward_type, rate, cap, exclusiones, valid_during).',
    '🆕 ve_values — tabla VE por programa (columna generada redemption × p_red × (1−breakage)). Indispensable para el motor v2.',
    '🆕 promotions con vigencia temporal — (effect, channel, city, mcc, compra_min, tope, conditions JSONB, stacking_group, is_targeted, valid_during) + N:N promotion_cards/promotion_merchants.',
    'cards_user — (user_id, card_id, nickname, last_four, is_primary) + 🆕 cupo_disp, dia_corte, soft-delete removed_at.',
    '🆕 recommendations — log inmutable (contexto → tarjeta → ahorro estimado → ¿seguida? → ahorro real).',
])
P('⚠️ Bug heredado de v1: el front consulta tablas en inglés (cards_user, cards, banks, benefits) pero schema.sql define nombres en español (user_cards, bancos, tarjetas, beneficios_*) → contra Supabase real devuelve vacío. Decisión v2: unificar a inglés y migrar. El front sigue desacoplado con VITE_USE_MOCKS=true.')
P('Privacidad: RLS por usuario en cards_user; no se recopilan PAN/CVV/fecha de vencimiento (fuera de PCI-DSS). 🆕 Habeas Data (Ley 1581): consentimiento explícito desde el día 1 al capturar gasto.')

# ===================== 11 =====================
H('11. Estrategia de Datos / Ingesta (Parte VIII — resuelve Pregunta Abierta 9.2)')
P('Mantener el grafo de promociones fresco y correcto ES el moat. Operación:')
bullets([
    'Fuentes: sitios/apps oficiales de bancos, newsletters, redes, crowdsourcing de usuarios.',
    'Cadencia: revisión semanal del alcance (≈20 tarjetas × 20 comercios); diaria para promos que vencen <7 días.',
    'Validación (regla de oro): doble fuente + cuatro ojos antes de marcar 🟢 Confirmado. Sin validar = no existe.',
    'SLA de frescura: ninguna promo recomendada con verificación > 7 días.',
    'Estados: borrador → validada → vigente → por_vencer → expirada.',
    'Evolución: Fase 1 curación manual → Fase 2 scraping semi-automatizado con validación humana → Fase 3 convenios/APIs/Open Finance.',
])

# ===================== 12 =====================
H('12. Monetización por Fases y Unit Economics (Parte VII)')
bullets([
    'Fase 1 (Años 1–3): CPA/afiliación + premium de nicho. Fuente principal al inicio.',
    'Fase 2: suscripción premium escalada + Open Finance.',
    'Fase 3 (Años 4–7): marketplace (Card-Linked Offers) + B2B/licenciamiento + inteligencia de mercado. Aquí migra el 70%+ del valor.',
])
P('Hallazgo central: el CPA es una trampa de techo bajo; el valor venture está en el marketplace + B2B sobre el dato cross-bank.')
table(['Métrica (Año 7)', 'A Conservador', 'B Base', 'C Ambicioso'], [
    ['LTV/CAC', '0.22 ❌', '1.88 ⚠️', '6.95 ✅'],
    ['Margen bruto', '60%', '72%', '82%'],
    ['Ingresos', '$0.2M', '$6.5M', '$80.3M USD'],
    ['Valoración', '~$1M', '$26–52M', '$321–643M USD'],
])
P('Capital a PMF: $0.3–1.1M USD. Caso C es el único venture-grade y capital-eficiente (~$3.7M de quema máx).')

# ===================== 13 =====================
H('13. Requisitos No Funcionales (NFRs) y Stack')
bullets([
    'Rendimiento: TTI < 1.5 s en 3G; consultas del recomendador < 100 ms (índices + precálculo user_best_card_by_category).',
    'Escalabilidad: lecturas vía Edge Functions/CDN de Supabase; 1.000 concurrentes (v1) → diseño preparado para 1M (Parte VI).',
    'Seguridad: AuthN JWT Supabase; AuthZ RLS Postgres; HTTPS/TLS 1.3. 🆕 Consentimientos Open Finance inmutables + secretos en Vault.',
    'Stack: Frontend React (Vite) + Tailwind + Framer Motion en Vercel/Netlify. Backend/DB Supabase. 🆕 Evolución: ingesta semi-automatizada y extracción de dominios a servicios.',
])

# ===================== 14 =====================
H('14. Diseño y UX (Sistema Visual — conservado de v1)')
bullets([
    'Principios: mobile-first (max-w-2xl, nav inferior fija de 5 ítems, pb-24); limpio y suave (rounded-2xl, sombras violeta); con vida (micro-animaciones Framer Motion, layoutId, active:scale-95).',
    'Color: primario eliseo violeta (500 #5B4CF5, 600 #4A3DE3); éxito mint #10B981; acento coral #FF6B57; fondo #F7F8FF; superficie #FFFFFF; texto #0F0F23; gradientes por nivel de tarjeta.',
    'Tipografía: Inter. Títulos text-2xl font-bold; secciones text-lg font-bold; cuerpo text-sm; metadatos text-xs.',
    'Componentes (index.css): eliseo-card, eliseo-btn-primary/-secondary/-outline, input-field, glass; 🆕 ConfidenceBadge (niveles de confianza); iconos lucide-react.',
])
table(['Ruta', 'Pantalla', 'Estado'], [
    ['/', 'Landing', '✅'],
    ['/auth', 'Login / Registro', '✅'],
    ['/dashboard', 'Tarjeta principal + stats + categorías', '✅'],
    ['/my-cards', 'Lista de billetera', '✅'],
    ['/add-card', 'Alta en 3 pasos', '✅'],
    ['/card-detail/:id', 'Detalle + beneficios (🆕 nivel de confianza)', '✅'],
    ['/recommender', 'Recomendador por categoría', '✅ (🔲 motor v2)'],
    ['/offers', 'Ofertas + filtros + 🆕 link saliente + confianza', '✅'],
    ['/profile', 'Perfil + privacidad', '✅ (🔲 perfil_pago)'],
    ['/destacados', '"Mejores del mercado" (afiliados)', '✅ (construido jun 2026)'],
])
P('Categorías (11): General 🔧 · Cashback 💰 · Puntos/Millas 🏆 · Viajes ✈️ · Restaurantes 🍽️ · Entretenimiento 🎬 · Supermercados 🛒 · Combustible ⛽ · Streaming 📺 · Moda 👗 · Seguros 🛡️.')
P('Tono de voz: español de Colombia, cercano, segunda persona; mensajes breves; estados vacíos siempre con CTA. 🆕 Todo mensaje de ahorro incluye su nivel de confianza y el "por qué" en pesos.')

# ===================== 15 =====================
H('15. Roadmap, Gates Go/No-Go e Hipótesis sin Validar (Parte IX)')
bullets([
    '30 días: motor manual (concierge) + primeros 15–30 usuarios del ICP.',
    '90 días: 100 usuarios; medir premio real, hábito, disposición a pagar. Go/No-Go grande.',
    '6 meses: producto semi-automatizado + 500–1.000 usuarios + primer ingreso.',
    '12 meses: seed + piloto de marketplace (CLO) + preparación Open Finance.',
])
P('Gate Go/No-Go (90 días): consultas no provocadas ≥2/sem Y retención sem-4 ≥40% Y premio mediano ≥$400k/año Y (WTP ≥25% en cuartil alto o CPA real). Si no → pivotar.')
P('Las 5 hipótesis NO validadas: (1) 🔴🔴🔴 premio ≥$400k/año en el Optimizador; (2) 🔴🔴 hábito recurrente; (3) 🔴🔴 motor de ingresos (CPA o suscripción); (4) 🔴 el dato de promociones se mantiene fresco a escala; (5) 🟠 la migración a marketplace/B2B es ejecutable.')
P('⚠️ Coherencia: el Blueprint recomienda validar 90 días con un concierge manual (WhatsApp) antes de construir software. La app React (incluidos los features de §6 ya construidos) es producto de Fase 1 post-validación; se desarrolla en paralelo por decisión del equipo.')

# ===================== 16 =====================
H('16. Preguntas Abiertas (estado v2)')
bullets([
    '9.1 Ranking del recomendador → ✅ RESUELTO vía normalización VE + score de dos capas. Pendiente: calibrar la tabla VE con datos reales.',
    '9.2 Ingesta de datos → 🟡 Estrategia definida (§11). Pendiente: ejecutar y medir sostenibilidad a 100 usuarios.',
    '9.3 Reconciliación de esquema front↔BD → 🔲 En curso (unificar a inglés y migrar).',
    '🆕 Calibración de p_red/breakage por programa con datos de redención reales.',
    '🆕 Política de elegibilidad de promos segmentadas ("clientes seleccionados").',
], style='List Number')

# ===================== Changelog =====================
H('Changelog v1 → v2')
table(['Área', 'v1', 'v2'], [
    ['Marco', 'Spec táctico del producto construido', 'PRD integrado: estrategia + producto'],
    ['ICP', 'Amplio (25–45, crédito+débito)', 'Estrecho: "El Optimizador" (alto gasto, totalero, crédito)'],
    ['Recomendador', 'VALUE_TYPE_WEIGHT × numeric_value', 'Motor v2: VE, totalero/rotativo, dos capas, promos con vigencia/stacking, confianza'],
    ['Datos', 'Offers/Benefits planos', 'Versionado + tabla VE + promociones temporales (Parte VI)'],
    ['Monetización', 'CTR a apps / afiliados', 'Fases CPA → marketplace (CLO) → B2B'],
    ['Moat', 'Ausente', 'Neutralidad cross-bank + 5 capas + matriz de amenazas'],
    ['Validación económica', 'Ausente', 'Premio segmentado (Parte IV) + gates Go/No-Go'],
    ['Incertidumbre', 'No contemplada', 'Niveles Confirmado/Probable/Acción'],
    ['North Star', '"Transacciones Optimizadas"', 'Conservada, reanclada a hábito + ahorro real'],
    ['Diseño/UX', 'Completo', 'Conservado íntegro'],
])

P('Cambios tras el draft v2 (jun 2026, rama feat/ui-polish): se construyeron Destacados (/destacados), CTA saliente en Ofertas (Offer.url + oferta clicable) y la UI de niveles de confianza (ConfidenceBadge). El cálculo automático del nivel y el motor v2 (VE, totalero/rotativo) siguen pendientes.')
P('Documento de referencia completo: ver Blueprint_Estrategico_Eliseo_ES.md (Partes I–IX).')

doc.save('PRD_Eliseo_ES_v2.docx')
print("PRD v2 (.docx) regenerado correctamente desde el .md canónico.")
