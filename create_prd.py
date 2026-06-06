from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Product Requirements Document (PRD) - Eliseo (Technical PM Edition)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Document status: aligned with code on feat/ui-polish (June 2026). Legend: ✅ Done · 🟡 Partial · 🔲 Pending.')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Executive Summary & KPIs', level=1)
doc.add_paragraph('Vision: Eliseo empowers Colombian consumers to effortlessly maximize the value of their credit and debit cards by providing real-time, category-based card recommendations, centralized benefits, and active offer discovery.', style='List Bullet')
doc.add_paragraph('Target Audience: Banked Colombians (25-45 y/o) who hold multiple payment cards and actively travel, dine out, or shop, but lack the time to track complex loyalty programs and benefits.', style='List Bullet')
doc.add_paragraph('Success Metrics (KPIs):', style='List Bullet')
doc.add_paragraph('Adoption: Average of >2.5 cards added per user during their first week.', style='List Number 2')
doc.add_paragraph('Engagement: >3 sessions per week per active user accessing the "Recommender" or "Offers" tab.', style='List Number 2')
doc.add_paragraph('Conversion (Monetization Proxy): >5% click-through rate (CTR) to official bank application links from the "Destacados" or "Offers" sections.', style='List Number 2')
doc.add_paragraph('North Star Metric: Number of "Optimized Transactions" generated (measured by how many times a user checks the recommender before making a purchase).', style='List Bullet')

doc.add_heading('2. Scope Definition (MoSCoW)', level=1)
doc.add_paragraph('Must Have (MVP - The "Walking Skeleton"):', style='List Bullet')
doc.add_paragraph('✅ User authentication (Sign up / Login).', style='List Number 2')
doc.add_paragraph('✅ Manual mechanism to add a card without storing sensitive data (Bank/model selection only, plus optional nickname and last 4 digits).', style='List Number 2')
doc.add_paragraph('✅ Dashboard showing primary card and summary metrics.', style='List Number 2')
doc.add_paragraph('✅ Category-based Recommender engine (matching user\'s added cards against the benefits DB).', style='List Number 2')
doc.add_paragraph('✅ Card Detail view showing categorized benefits.', style='List Number 2')

doc.add_paragraph('Should Have:', style='List Bullet')
doc.add_paragraph('✅ Dynamic Offers feed with expiration dates.', style='List Number 2')
doc.add_paragraph('✅ Filtering of Offers by Bank and Category.', style='List Number 2')
doc.add_paragraph('🔲 Highlighting of "Best in Market" cards (Destacados) to drive affiliate conversions. (No screen or route yet.)', style='List Number 2')

doc.add_paragraph('Could Have:', style='List Bullet')
doc.add_paragraph('🔲 Push notifications for expiring offers. (Existed in previous main; not in current refactor.)', style='List Number 2')
doc.add_paragraph('🔲 Geolocated benefits (e.g., "You are near an airport, here are your Lounge benefits"). (Same.)', style='List Number 2')

doc.add_paragraph("Won't Have (v1.0):", style='List Bullet')
doc.add_paragraph('Direct API integration with banks for real-time transaction history tracking (Open Banking).', style='List Number 2')
doc.add_paragraph('Automatic points/miles balance scraping.', style='List Number 2')

doc.add_heading('3. User Personas & Jobs to be Done (JTBD)', level=1)
doc.add_paragraph('Note: detailed, current personas live in .product/PERSONAS/ (Gloria Espinosa, Sandra Moreno, Sebastián Ríos). The "Andrés" archetype below is kept as a summary of the primary JTBD.', style='List Bullet')
doc.add_paragraph('Primary archetype (JTBD): Andrés (32), Tech-savvy Millennial. Works in tech/marketing. Has 4 cards (Nu, Bancolombia LifeMiles, RappiCard, Falabella). Travels 3 times a year. Wants to get free flights but forgets which card gives lounge access or the best miles multiplier.', style='List Bullet')
doc.add_paragraph('JTBD:', style='List Bullet')
doc.add_paragraph('"When I am waiting to pay at a restaurant, I want to quickly see which of my 4 cards gives me the best cashback, so that I don\'t leave money on the table."', style='List Number 2')
doc.add_paragraph('"When I am planning a trip, I want to know which of my cards offers free airport lounge access, so that I can wait comfortably without paying $30 USD."', style='List Number 2')

doc.add_heading('4. Functional Requirements (User Stories)', level=1)

doc.add_heading('4.1. Core: Manage Wallet', level=2)
doc.add_paragraph('Story: As a user, I want to add a card to my digital wallet by selecting the bank and the card model, so that the app knows what benefits I have.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('[x] System displays a searchable list of banks (traditional and digital).', style='List Number 2')
doc.add_paragraph('[x] System filters card models immediately after bank selection.', style='List Number 2')
doc.add_paragraph('[x] User can optionally enter a string (max 40 chars) for Nickname, and numeric (max 4 chars) for "Last 4 digits".', style='List Number 2')
doc.add_paragraph('[x] System saves the mapping (user_id, card_id, nickname, last_four) to the database.', style='List Number 2')

doc.add_heading('4.2. Core: Recommendation Engine', level=2)
doc.add_paragraph('Story: As a user, I want to select a spending category (e.g., "Restaurants") so that the app ranks my cards from best to worst for that purchase.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('[x] System retrieves all cards linked to the active user.', style='List Number 2')
doc.add_paragraph('[x] System fetches benefits for those cards matching the selected category_id.', style='List Number 2')
doc.add_paragraph('[x] System sorts cards based on a weighting algorithm. (Implemented in Recommender.tsx via VALUE_TYPE_WEIGHT × numeric_value; see open question 9.1.)', style='List Number 2')
doc.add_paragraph('[x] Results load in < 500ms.', style='List Number 2')
doc.add_paragraph('[x] If no card has specific benefits for the category, system displays a fallback message teaching the user to use their baseline card.', style='List Number 2')

doc.add_heading('4.3. Core: Explorable Offers', level=2)
doc.add_paragraph('Story: As a user, I want to filter current market offers by category or bank, so that I can find discounts I can use this weekend.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('[x] System queries the Offers table for active promotions (valid_until >= today OR valid_until is NULL).', style='List Number 2')
doc.add_paragraph('[x] User can toggle one or multiple bank chips to filter.', style='List Number 2')
doc.add_paragraph('[x] User can toggle category chips to filter.', style='List Number 2')
doc.add_paragraph('[ ] 🔲 Pending: Clicking an offer opens a detailed view with a direct, trackable outbound link to the bank. (Today offers are not clickable and the Offer type has no URL field. This is the CTA the conversion KPI depends on.)', style='List Number 2')

doc.add_heading('5. Core Logical Flows', level=1)

doc.add_heading('Flow 1: Add a Card (Happy Path)', level=2)
doc.add_paragraph('1. User clicks "+" or "Agregar tarjeta".', style='List Number')
doc.add_paragraph('2. System fetches SELECT id, name, logo FROM banks ORDER BY name. User selects bank "Bancolombia".', style='List Number')
doc.add_paragraph('3. System fetches SELECT * FROM cards WHERE bank_id = "Bancolombia". User selects "Visa LifeMiles Platinum".', style='List Number')
doc.add_paragraph('4. User enters nickname "Mi tarjeta principal".', style='List Number')
doc.add_paragraph('5. User clicks "Guardar".', style='List Number')
doc.add_paragraph('6. Frontend sends POST /api/user-cards with payload { user_id: "uuid", card_id: "uuid", nickname: "Mi tarjeta principal", is_primary: true/false }.', style='List Number')
doc.add_paragraph('7. Backend validates payload and inserts into cards_user. Returns 200 OK.', style='List Number')
doc.add_paragraph('8. Frontend redirects to Dashboard /dashboard and displays the newly added card via optimistic UI update.', style='List Number')

doc.add_heading('Flow 2: The Recommender (Edge Case: No Cards Added)', level=2)
doc.add_paragraph('1. User navigates to /recommender and selects "Combustible".', style='List Number')
doc.add_paragraph('2. System checks user\'s wallet. Wallet count = 0.', style='List Number')
doc.add_paragraph('3. System interrupts flow and shows Empty State: "Agrega tu primera tarjeta para obtener recomendaciones personalizadas."', style='List Number')
doc.add_paragraph('4. User clicks CTA -> Redirects to Add Card flow.', style='List Number')

doc.add_heading('6. Data Requirements', level=1)
doc.add_paragraph('Key Entities:', style='List Bullet')
doc.add_paragraph('Users: id (UUID), email, created_at. (Managed via Supabase Auth).', style='List Number 2')
doc.add_paragraph('Banks: id (string/PK), name, short_name, is_digital, logo_color, website.', style='List Number 2')
doc.add_paragraph('Cards: id (PK), bank_id (FK), name, franchise (Visa/MC/Amex), tier (Classic, Gold, Black...), type (Credit/Debit).', style='List Number 2')
doc.add_paragraph('Benefits: id (PK), card_id (FK), category (enum: viajes, restaurantes, etc.), value_type (cashback_percent, points_multiplier, etc.), value (numeric).', style='List Number 2')
doc.add_paragraph('Cards_User (Junction): id (PK), user_id (FK), card_id (FK), nickname, last_four, is_primary (boolean).', style='List Number 2')
doc.add_paragraph('Offers: id (PK), bank_id (FK), title, description, category, valid_until (timestamp). (Pending: add url/outbound link for the 4.3 CTA.)', style='List Number 2')
doc.add_paragraph('⚠️ Front ↔ DB mismatch to reconcile: the front queries English table names (cards_user, cards, banks, benefits), but supabase/schema.sql defines Spanish names (user_cards, bancos, tarjetas, beneficios_*). Until unified, the front against real Supabase returns empty data. The type contract lives in src/types/database.ts. The front can be developed decoupled with mock data (VITE_USE_MOCKS=true).', style='List Bullet')

doc.add_paragraph('Privacy:', style='List Bullet')
doc.add_paragraph('Data Scoping: Row Level Security (RLS) policies must ensure users can ONLY read/write their own cards_user records.', style='List Number 2')
doc.add_paragraph('Sensitive Data: Real credit card PANs, expiration dates, and CVVs are explicitly NOT collected or processed to avoid PCI-DSS scope entirely.', style='List Number 2')

doc.add_heading('7. Non-Functional Requirements (NFRs)', level=1)
doc.add_paragraph('Performance:', style='List Bullet')
doc.add_paragraph('Time To Interactive (TTI) < 1.5 seconds on 3G mobile networks.', style='List Number 2')
doc.add_paragraph('DB Queries for the Recommender must execute in < 100ms.', style='List Number 2')
doc.add_paragraph('Scalability: Horizontal scaling for reads via Supabase Edge Functions/CDN, supporting up to 1,000 Concurrent Users.', style='List Bullet')
doc.add_paragraph('Security: AuthN via Supabase JWT. AuthZ managed via RLS in Postgres. HTTPS/TLS 1.3 forced on all connections.', style='List Bullet')
doc.add_paragraph('Maintainability: Strict TypeScript interfaces for all DB entity mappings. Standardized error handling mechanism.', style='List Bullet')

doc.add_heading('8. Technical Constraints & Stack Recommendation', level=1)
doc.add_paragraph('Constraints: Launch must be completed rapidly to test market fit. Budget is limited, requiring free-tier or pay-as-you-go infrastructure.', style='List Bullet')
doc.add_paragraph('Recommended Stack:', style='List Bullet')
doc.add_paragraph('Frontend: React (Vite) + Tailwind CSS + Framer Motion (for glassmorphism and Micro-animations). Hosted on Vercel or Netlify.', style='List Number 2')
doc.add_paragraph('Backend/DB: Supabase (PostgreSQL, Auth, Edge Functions). Eliminates the need to write boilerplate CRUD APIs and offers out-of-the-box realtime capabilities and RLS.', style='List Number 2')

doc.add_heading('9. Open Questions / Blockers', level=1)
doc.add_paragraph('1. Recommender Algorithm Logic: How do we objectively rank a card offering "10% cashback" vs a card offering "x3 LATAM Miles"? Does the user need a "Preferences" screen to weight miles vs. cashback?', style='List Number')
doc.add_paragraph('2. Data Ingestion: Manual entry of Card Benefits does not scale. What is the process or tool for scraping bank websites to keep the Offers table and Benefits table updated?', style='List Number')
doc.add_paragraph('3. Schema reconciliation: unify table names between front (cards_user/banks/benefits) and DB (user_cards/bancos/beneficios_*) — rename the schema or adjust the front? (In progress.)', style='List Number')

doc.add_heading('10. Design & UX (Visual System)', level=1)
doc.add_paragraph('Section added as a base for the redesign. Reflects the current system in tailwind.config.js + src/index.css.')

doc.add_heading('10.1 Principles', level=2)
doc.add_paragraph('Mobile-first: centered container max-w-2xl, fixed bottom navigation of 5 items (Home, Cards, Recommender, Offers, Profile). Content reserves pb-24 so it is not covered by the nav.', style='List Bullet')
doc.add_paragraph('Clean and soft: rounded cards (rounded-2xl), subtle violet-tinted shadows, plenty of whitespace, hierarchy via type weight.', style='List Bullet')
doc.add_paragraph('Alive: entry micro-animations (fade + slide) with Framer Motion; animated nav indicator (layoutId), active:scale-95 on buttons.', style='List Bullet')

doc.add_heading('10.2 Color', level=2)
doc.add_paragraph('Primary eliseo (violet): 50 #F0EFFE · 100 #E4E1FD · 400 #8577F5 · 500 #5B4CF5 (primary action) · 600 #4A3DE3.', style='List Bullet')
doc.add_paragraph('Success mint: 50 #EAFBF5 · 500 #10B981.', style='List Bullet')
doc.add_paragraph('Accent coral 500 #FF6B57 (alerts/highlights).', style='List Bullet')
doc.add_paragraph('Background #F7F8FF · Surface #FFFFFF · Text #0F0F23.', style='List Bullet')
doc.add_paragraph('Gradients: gradient-eliseo (brand), and per card tier: gradient-gold, gradient-platinum, gradient-black, etc.', style='List Bullet')

doc.add_heading('10.3 Typography', level=2)
doc.add_paragraph('Inter font (system-ui fallback). Page titles text-2xl font-bold; sections section-title (text-lg font-bold); body text-sm; metadata text-xs / text-[10-11px].', style='List Bullet')

doc.add_heading('10.4 Components (classes in index.css)', level=2)
doc.add_paragraph('eliseo-card — white container, rounded-2xl, card shadow, eliseo-100/50 border.', style='List Bullet')
doc.add_paragraph('eliseo-btn-primary / -secondary / -outline — rounded-xl buttons, active:scale-95.', style='List Bullet')
doc.add_paragraph('input-field — input with focus:ring-eliseo-300.', style='List Bullet')
doc.add_paragraph('glass — bottom bar with backdrop-blur. Shadows: card, card-hover, float. Icons: lucide-react.', style='List Bullet')

doc.add_heading('10.5 Screen inventory', level=2)
screens = [
    ('/', 'Landing (hero + features)', '✅'),
    ('/auth', 'Login / Sign up', '✅'),
    ('/dashboard', 'Primary card + stats + quick actions + categories', '✅'),
    ('/my-cards', 'List of wallet cards', '✅'),
    ('/add-card', 'Add in 3 steps (bank → card → details)', '✅'),
    ('/card-detail/:id', 'Detail + benefits by category + remove/primary', '✅'),
    ('/recommender', 'Recommender by category', '✅'),
    ('/offers', 'Offers feed with filters', '🟡 (missing detail + outbound link)'),
    ('/profile', 'Profile + privacy + sign out', '✅'),
    ('/destacados', '"Best in market"', '🔲 (to be designed)'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Route', 'Screen', 'Status'
for route, screen, status in screens:
    row = table.add_row().cells
    row[0].text, row[1].text, row[2].text = route, screen, status

doc.add_heading('10.6 Categories (11, defined in src/types/database.ts)', level=2)
doc.add_paragraph('General 🔧 · Cashback 💰 · Points/Miles 🏆 · Travel ✈️ · Restaurants 🍽️ · Entertainment 🎬 · Supermarkets 🛒 · Fuel ⛽ · Streaming 📺 · Fashion 👗 · Insurance 🛡️.')

doc.add_heading('10.7 Tone of voice', level=2)
doc.add_paragraph('Colombian Spanish, warm and clear, second person ("tus tarjetas", "elige una categoría"). Short messages; empty states always with a CTA. Correct accentuation required.')

doc.save('PRD_Eliseo.docx')
print("English document generated successfully.")
