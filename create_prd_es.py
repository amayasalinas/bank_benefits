from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Documento de Requisitos del Producto (PRD) - Eliseo (Edición de PM Técnico)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Estado del documento: ajustado al código en feat/ui-polish (junio 2026). Leyenda: ✅ Hecho · 🟡 Parcial · 🔲 Pendiente.')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Resumen Ejecutivo y KPIs', level=1)
doc.add_paragraph('Visión: Eliseo empodera a los consumidores colombianos para maximizar de manera sencilla el valor de sus tarjetas de crédito y débito proporcionando recomendaciones de tarjetas en tiempo real basadas en categorías, beneficios centralizados y descubrimiento activo de ofertas.', style='List Bullet')
doc.add_paragraph('Público Objetivo: Colombianos bancarizados (25-45 años) que poseen múltiples tarjetas de pago y viajan, comen fuera o compran activamente, pero carecen de tiempo para hacer seguimiento a programas de lealtad y beneficios complejos.', style='List Bullet')
doc.add_paragraph('Métricas de Éxito (KPIs):', style='List Bullet')
doc.add_paragraph('Adopción: Promedio de >2.5 tarjetas agregadas por usuario durante su primera semana.', style='List Number 2')
doc.add_paragraph('Engagement: >3 sesiones por semana por usuario activo accediendo a la pestaña de "Recomendador" u "Ofertas".', style='List Number 2')
doc.add_paragraph('Conversión (Proxy de Monetización): >5% de tasa de clics (CTR) hacia enlaces de aplicaciones bancarias oficiales desde las secciones de "Destacados" u "Ofertas".', style='List Number 2')
doc.add_paragraph('Métrica North Star: Número de "Transacciones Optimizadas" generadas (medido por cuántas veces un usuario revisa el recomendador antes de realizar una compra).', style='List Bullet')

doc.add_heading('2. Definición del Alcance (MoSCoW)', level=1)
doc.add_paragraph('Must Have (MVP - El "Esqueleto Andante"):', style='List Bullet')
doc.add_paragraph('✅ Autenticación de usuarios (Registro / Inicio de sesión).', style='List Number 2')
doc.add_paragraph('✅ Mecanismo manual para agregar una tarjeta sin almacenar datos sensibles (Solo selección de Banco/modelo, más apodo opcional y últimos 4 dígitos).', style='List Number 2')
doc.add_paragraph('✅ Panel de control (Dashboard) mostrando la tarjeta principal y métricas resumidas.', style='List Number 2')
doc.add_paragraph('✅ Motor Recomendador basado en categorías (emparejando las tarjetas agregadas del usuario contra la base de beneficios).', style='List Number 2')
doc.add_paragraph('✅ Vista de Detalle de Tarjeta mostrando beneficios categorizados.', style='List Number 2')

doc.add_paragraph('Should Have:', style='List Bullet')
doc.add_paragraph('✅ Feed dinámico de Ofertas con fechas de vencimiento.', style='List Number 2')
doc.add_paragraph('✅ Filtrado de Ofertas por Banco y Categoría.', style='List Number 2')
doc.add_paragraph('🔲 Destacar las tarjetas "Mejores del Mercado" (Destacados) para impulsar conversiones de afiliados. (No existe pantalla ni ruta todavía.)', style='List Number 2')

doc.add_paragraph('Could Have:', style='List Bullet')
doc.add_paragraph('🔲 Notificaciones push para ofertas a punto de expirar. (Existían en main anterior; no están en el refactor actual.)', style='List Number 2')
doc.add_paragraph('🔲 Beneficios geolocalizados (ej., "Estás cerca de un aeropuerto, aquí están tus beneficios de Salas VIP"). (Ídem.)', style='List Number 2')

doc.add_paragraph("Won't Have (v1.0):", style='List Bullet')
doc.add_paragraph('Integración directa de API con bancos para seguimiento del historial de transacciones en tiempo real (Open Banking).', style='List Number 2')
doc.add_paragraph('Extracción (scraping) automática del saldo de puntos/millas.', style='List Number 2')

doc.add_heading('3. Personas de Usuario y Tareas a Realizar (JTBD)', level=1)
doc.add_paragraph('Nota: las personas detalladas y vigentes viven en .product/PERSONAS/ (Gloria Espinosa, Sandra Moreno, Sebastián Ríos). El arquetipo "Andrés" se conserva como resumen del JTBD principal.', style='List Bullet')
doc.add_paragraph('Arquetipo principal (JTBD): Andrés (32), Millennial experto en tecnología. Trabaja en tecnología/marketing. Tiene 4 tarjetas (Nu, Bancolombia LifeMiles, RappiCard, Falabella). Viaja 3 veces al año. Quiere obtener vuelos gratis pero olvida qué tarjeta le da acceso a salas VIP o el mejor multiplicador de millas.', style='List Bullet')
doc.add_paragraph('JTBD:', style='List Bullet')
doc.add_paragraph('"Cuando estoy esperando para pagar en un restaurante, quiero ver rápidamente cuál de mis 4 tarjetas me da el mejor cashback, para no perder dinero."', style='List Number 2')
doc.add_paragraph('"Cuando estoy planeando un viaje, quiero saber cuál de mis tarjetas ofrece acceso gratuito a salas VIP de aeropuertos, para poder esperar cómodamente sin pagar $30 USD."', style='List Number 2')

doc.add_heading('4. Requisitos Funcionales (Historias de Usuario)', level=1)

doc.add_heading('4.1. Core: Gestionar Billetera', level=2)
doc.add_paragraph('Historia: Como usuario, quiero agregar una tarjeta a mi billetera digital seleccionando el banco y el modelo de la tarjeta, para que la aplicación sepa qué beneficios tengo.', style='List Bullet')
doc.add_paragraph('Criterios de Aceptación:', style='List Bullet')
doc.add_paragraph('[x] El sistema muestra una lista de bancos (tradicionales y digitales) donde se puede buscar.', style='List Number 2')
doc.add_paragraph('[x] El sistema filtra los modelos de tarjetas inmediatamente después de seleccionar el banco.', style='List Number 2')
doc.add_paragraph('[x] El usuario puede ingresar opcionalmente una cadena (máx. 40 caracteres) para el Apodo, y numérica (máx. 4 caracteres) para los "Últimos 4 dígitos".', style='List Number 2')
doc.add_paragraph('[x] El sistema guarda el mapeo (user_id, card_id, nickname, last_four) en la base de datos.', style='List Number 2')

doc.add_heading('4.2. Core: Motor de Recomendación', level=2)
doc.add_paragraph('Historia: Como usuario, quiero seleccionar una categoría de gasto (ej., "Restaurantes") para que la aplicación clasifique mis tarjetas de mejor a peor para esa compra.', style='List Bullet')
doc.add_paragraph('Criterios de Aceptación:', style='List Bullet')
doc.add_paragraph('[x] El sistema recupera todas las tarjetas vinculadas al usuario activo.', style='List Number 2')
doc.add_paragraph('[x] El sistema obtiene los beneficios de esas tarjetas que coinciden con el category_id seleccionado.', style='List Number 2')
doc.add_paragraph('[x] El sistema ordena las tarjetas basándose en un algoritmo de ponderación. (Implementado en Recommender.tsx vía VALUE_TYPE_WEIGHT × numeric_value; ver pregunta abierta 9.1.)', style='List Number 2')
doc.add_paragraph('[x] Los resultados cargan en < 500ms.', style='List Number 2')
doc.add_paragraph('[x] Si ninguna tarjeta tiene beneficios específicos para la categoría, el sistema muestra un mensaje de respaldo enseñando al usuario a usar su tarjeta base principal.', style='List Number 2')

doc.add_heading('4.3. Core: Ofertas Explorables', level=2)
doc.add_paragraph('Historia: Como usuario, quiero filtrar las ofertas actuales del mercado por categoría o banco, para poder encontrar descuentos que pueda usar este fin de semana.', style='List Bullet')
doc.add_paragraph('Criterios de Aceptación:', style='List Bullet')
doc.add_paragraph('[x] El sistema consulta la tabla de Offers (Ofertas) para promociones activas (valid_until >= hoy O valid_until es NULL).', style='List Number 2')
doc.add_paragraph('[x] El usuario puede alternar uno o múltiples chips/etiquetas de bancos para filtrar.', style='List Number 2')
doc.add_paragraph('[x] El usuario puede alternar chips/etiquetas de categorías para filtrar.', style='List Number 2')
doc.add_paragraph('[ ] 🔲 Pendiente: Hacer clic en una oferta abre una vista detallada con un enlace saliente directo y rastreable hacia el banco. (Hoy la oferta no es clicable y el tipo Offer no tiene campo URL. Es el CTA del que depende el KPI de conversión.)', style='List Number 2')

doc.add_heading('5. Flujos Lógicos Principales', level=1)

doc.add_heading('Flujo 1: Agregar una Tarjeta (Ruta Feliz)', level=2)
doc.add_paragraph('1. El usuario hace clic en "+" o "Agregar tarjeta".', style='List Number')
doc.add_paragraph('2. El sistema recupera SELECT id, name, logo FROM banks ORDER BY name. El usuario selecciona el banco "Bancolombia".', style='List Number')
doc.add_paragraph('3. El sistema recupera SELECT * FROM cards WHERE bank_id = "Bancolombia". El usuario selecciona "Visa LifeMiles Platinum".', style='List Number')
doc.add_paragraph('4. El usuario ingresa el apodo "Mi tarjeta principal".', style='List Number')
doc.add_paragraph('5. El usuario hace clic en "Guardar".', style='List Number')
doc.add_paragraph('6. El frontend envía POST /api/user-cards con el payload { user_id: "uuid", card_id: "uuid", nickname: "Mi tarjeta principal", is_primary: true/false }.', style='List Number')
doc.add_paragraph('7. El backend valida el payload e inserta en cards_user. Devuelve 200 OK.', style='List Number')
doc.add_paragraph('8. El frontend redirige al Dashboard /dashboard y muestra la tarjeta recién agregada mediante una actualización optimista de la UI.', style='List Number')

doc.add_heading('Flujo 2: El Recomendador (Caso Extremo: Sin Tarjetas Agregadas)', level=2)
doc.add_paragraph('1. El usuario navega a /recommender y selecciona "Combustible".', style='List Number')
doc.add_paragraph('2. El sistema comprueba la billetera del usuario. Recuento de billetera = 0.', style='List Number')
doc.add_paragraph('3. El sistema interrumpe el flujo y muestra un Estado Vacío: "Agrega tu primera tarjeta para obtener recomendaciones personalizadas."', style='List Number')
doc.add_paragraph('4. El usuario hace clic en el CTA (Llamado a la acción) -> Redirige al flujo de Agregar Tarjeta.', style='List Number')

doc.add_heading('6. Requisitos de Datos', level=1)
doc.add_paragraph('Entidades Clave:', style='List Bullet')
doc.add_paragraph('Usuarios (Users): id (UUID), email, created_at. (Manejado vía Supabase Auth).', style='List Number 2')
doc.add_paragraph('Bancos (Banks): id (cadena/PK), name, short_name, is_digital, logo_color, website.', style='List Number 2')
doc.add_paragraph('Tarjetas (Cards): id (PK), bank_id (FK), name, franchise (Visa/MC/Amex), tier (Classic, Gold, Black...), type (Credit/Debit).', style='List Number 2')
doc.add_paragraph('Beneficios (Benefits): id (PK), card_id (FK), category (enum: viajes, restaurantes, etc.), value_type (cashback_percent, points_multiplier, etc.), value (numérico).', style='List Number 2')
doc.add_paragraph('Tarjetas_Usuario (Unión / Cards_User): id (PK), user_id (FK), card_id (FK), nickname, last_four, is_primary (booleano).', style='List Number 2')
doc.add_paragraph('Ofertas (Offers): id (PK), bank_id (FK), title, description, category, valid_until (timestamp). (Pendiente agregar url/link saliente para el CTA de la 4.3.)', style='List Number 2')
doc.add_paragraph('⚠️ Desajuste front ↔ BD a reconciliar: el front consulta tablas en inglés (cards_user, cards, banks, benefits), pero supabase/schema.sql define nombres en español (user_cards, bancos, tarjetas, beneficios_*). Hasta unificar nombres, el front contra Supabase real devuelve datos vacíos. El contrato de tipos vive en src/types/database.ts. El front puede desarrollarse desacoplado con datos mock (VITE_USE_MOCKS=true).', style='List Bullet')

doc.add_paragraph('Privacidad:', style='List Bullet')
doc.add_paragraph('Alcance de Datos: Las políticas de Seguridad a Nivel de Fila (Row Level Security - RLS) deben garantizar que los usuarios SOLO puedan leer/escribir sus propios registros en cards_user.', style='List Number 2')
doc.add_paragraph('Datos Sensibles: Se establece explícitamente que los números de tarjeta (PAN), fechas de vencimiento y CVV reales NO se recopilan ni procesan para evitar por completo el alcance de la normativa PCI-DSS.', style='List Number 2')

doc.add_heading('7. Requisitos No Funcionales (NFRs)', level=1)
doc.add_paragraph('Rendimiento:', style='List Bullet')
doc.add_paragraph('Tiempo para ser Interactivo (TTI) < 1.5 segundos en redes móviles 3G.', style='List Number 2')
doc.add_paragraph('Las consultas a la base de datos para el Recomendador deben ejecutarse en < 100ms.', style='List Number 2')
doc.add_paragraph('Escalabilidad: Escalado horizontal para lecturas a través de las Edge Functions/CDN de Supabase, soportando hasta 1,000 Usuarios Concurrentes.', style='List Bullet')
doc.add_paragraph('Seguridad: Autenticación (AuthN) a través de JWT de Supabase. Autorización (AuthZ) administrada mediante RLS en Postgres. HTTPS/TLS 1.3 forzado en todas las conexiones.', style='List Bullet')
doc.add_paragraph('Mantenibilidad: Interfaces estrictas de TypeScript para todos los mapeos de entidades de la base de datos. Mecanismo estandarizado de manejo de errores.', style='List Bullet')

doc.add_heading('8. Restricciones Técnicas y Recomendación de Stack Tecnológico', level=1)
doc.add_paragraph('Restricciones: El lanzamiento debe completarse rápidamente para probar el ajuste en el mercado (market fit). El presupuesto es limitado, requiriendo infraestructura de nivel gratuito (free-tier) o pago por uso.', style='List Bullet')
doc.add_paragraph('Stack Recomendado:', style='List Bullet')
doc.add_paragraph('Frontend: React (Vite) + Tailwind CSS + Framer Motion (para glassmorphism y micro-animaciones). Alojado en Vercel o Netlify.', style='List Number 2')
doc.add_paragraph('Backend/DB: Supabase (PostgreSQL, Auth, Edge Functions). Elimina la necesidad de escribir APIs CRUD repetitivas (boilerplate) y ofrece capacidades en tiempo real y RLS desde el primer momento.', style='List Number 2')

doc.add_heading('9. Preguntas Abiertas / Bloqueos', level=1)
doc.add_paragraph('1. Lógica del Algoritmo Recomendador: ¿Cómo clasificamos objetivamente una tarjeta que ofrece "10% de cashback" frente a una que ofrece "x3 Millas LATAM"? ¿El usuario necesita una pantalla de "Preferencias" para ponderar las millas frente al cashback?', style='List Number')
doc.add_paragraph('2. Ingesta de Datos: El ingreso manual de los Beneficios de las Tarjetas no es escalable. ¿Cuál es el proceso o herramienta para extraer datos (scraping) de los sitios web de los bancos y mantener actualizadas las tablas de Offers y Benefits?', style='List Number')
doc.add_paragraph('3. Reconciliación de esquema: unificar los nombres de tabla entre front (cards_user/banks/benefits) y BD (user_cards/bancos/beneficios_*) — ¿se renombra el esquema o se ajusta el front? (En curso.)', style='List Number')

doc.add_heading('10. Diseño y UX (Sistema Visual)', level=1)
doc.add_paragraph('Sección añadida como base para el rediseño. Refleja el sistema actual en tailwind.config.js + src/index.css.')

doc.add_heading('10.1 Principios', level=2)
doc.add_paragraph('Mobile-first: contenedor centrado max-w-2xl, navegación inferior fija de 5 ítems (Inicio, Tarjetas, Recomendador, Ofertas, Perfil). El contenido reserva pb-24 para no quedar tapado por el nav.', style='List Bullet')
doc.add_paragraph('Limpio y suave: tarjetas redondeadas (rounded-2xl), sombras tenues de tinte violeta, mucho aire en blanco, jerarquía por peso tipográfico.', style='List Bullet')
doc.add_paragraph('Con vida: micro-animaciones de entrada (fade + slide) con Framer Motion; indicador de nav animado (layoutId), active:scale-95 en botones.', style='List Bullet')

doc.add_heading('10.2 Color', level=2)
doc.add_paragraph('Primario eliseo (violeta): 50 #F0EFFE · 100 #E4E1FD · 400 #8577F5 · 500 #5B4CF5 (acción principal) · 600 #4A3DE3.', style='List Bullet')
doc.add_paragraph('Éxito mint: 50 #EAFBF5 · 500 #10B981.', style='List Bullet')
doc.add_paragraph('Acento coral 500 #FF6B57 (alertas/realces).', style='List Bullet')
doc.add_paragraph('Fondo background #F7F8FF · Superficie surface #FFFFFF · Texto #0F0F23.', style='List Bullet')
doc.add_paragraph('Gradientes: gradient-eliseo (marca), y por nivel de tarjeta: gradient-gold, gradient-platinum, gradient-black, etc.', style='List Bullet')

doc.add_heading('10.3 Tipografía', level=2)
doc.add_paragraph('Fuente Inter (system-ui fallback). Títulos de página text-2xl font-bold; secciones section-title (text-lg font-bold); cuerpo text-sm; metadatos text-xs / text-[10-11px].', style='List Bullet')

doc.add_heading('10.4 Componentes (clases en index.css)', level=2)
doc.add_paragraph('eliseo-card — contenedor blanco, rounded-2xl, sombra card, borde eliseo-100/50.', style='List Bullet')
doc.add_paragraph('eliseo-btn-primary / -secondary / -outline — botones rounded-xl, active:scale-95.', style='List Bullet')
doc.add_paragraph('input-field — input con focus:ring-eliseo-300.', style='List Bullet')
doc.add_paragraph('glass — barra inferior con backdrop-blur. Sombras: card, card-hover, float. Iconos: lucide-react.', style='List Bullet')

doc.add_heading('10.5 Inventario de pantallas', level=2)
screens = [
    ('/', 'Landing (hero + features)', '✅'),
    ('/auth', 'Login / Registro', '✅'),
    ('/dashboard', 'Tarjeta principal + stats + accesos rápidos + categorías', '✅'),
    ('/my-cards', 'Lista de tarjetas de la cartera', '✅'),
    ('/add-card', 'Alta en 3 pasos (banco → tarjeta → detalles)', '✅'),
    ('/card-detail/:id', 'Detalle + beneficios por categoría + eliminar/principal', '✅'),
    ('/recommender', 'Recomendador por categoría', '✅'),
    ('/offers', 'Feed de ofertas con filtros', '🟡 (falta detalle + link saliente)'),
    ('/profile', 'Perfil + privacidad + cerrar sesión', '✅'),
    ('/destacados', '"Mejores del mercado"', '🔲 (por diseñar)'),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = 'Ruta', 'Pantalla', 'Estado'
for ruta, pantalla, estado in screens:
    row = table.add_row().cells
    row[0].text, row[1].text, row[2].text = ruta, pantalla, estado

doc.add_heading('10.6 Categorías (11, definidas en src/types/database.ts)', level=2)
doc.add_paragraph('General 🔧 · Cashback 💰 · Puntos/Millas 🏆 · Viajes ✈️ · Restaurantes 🍽️ · Entretenimiento 🎬 · Supermercados 🛒 · Combustible ⛽ · Streaming 📺 · Moda 👗 · Seguros 🛡️.')

doc.add_heading('10.7 Tono de voz', level=2)
doc.add_paragraph('Español de Colombia, cercano y claro, en segunda persona ("tus tarjetas", "elige una categoría"). Mensajes breves; estados vacíos siempre con un CTA. Acentuación correcta obligatoria.')

doc.save('PRD_Eliseo_ES.docx')
print("Documento ES generado correctamente.")
